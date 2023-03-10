import mysql.connector
import os
import matplotlib.pyplot as plt
import networkx as nx
from docx import Document
import numpy as np
from matplotlib.ticker import MaxNLocator
import matplotlib
import math
from sklearn.cluster import KMeans
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from docx.shared import Inches
import io
import base64
from docxcompose.composer import Composer
import scipy.stats as st
from pyecharts.charts import Geo
from pyecharts.charts import Map
from sklearn.metrics import r2_score
from docx2pdf import convert
import pandas as pd
import seaborn as sns
from matplotlib.ticker import MaxNLocator
from pyecharts import options as opts
from pyecharts.charts import Line,Grid,Bar,Pie,Radar
from pyecharts.faker import Faker
from pyecharts.globals import ThemeType
from pyecharts.globals import ChartType
import json
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from pyecharts.render import make_snapshot
import docx
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_UNDERLINE
from docx.enum.text import WD_TAB_ALIGNMENT
from docx.enum.table import WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_LINE_SPACING
from snapshot_selenium import snapshot
import yaml

class mydb:#数据库操作类

    def __init__(self,host,user,password,database,port,charset):
        self.host=host
        self.user=user
        self.password=password
        self.database=database
        self.port=port
        self.charset=charset
        
    def myconnect(self):#连接数据库
        conn = mysql.connector.connect(
            host=self.host,
            user=self.user,
            password=self.password,
            database=self.database,
            port=self.port,
            charset = self.charset
            # auth_plugin='mysql_native_password'
        )
        return conn

    def myexecu(self,execuword):#执行查询语句并保存在data中
        self.conn=self.myconnect()
        cursor=self.conn.cursor()
        cursor.execute(execuword)        
        self.data = cursor.fetchall()
        self.conn.commit()
        cursor.close()


    def endconn(self):#断开数据库连接
        self.conn.close()

class cladata:#基类

    def __init__(self,db,title):#将data内容传递给cladata类
        self.data=db.data
        self.title=title
        self.db=db


class authoranaly(cladata):#申请人分析

    def authorpie(self,dic,title):#绘制饼状图
        c =Pie(init_opts=opts.InitOpts(theme=ThemeType.WALDEN))
        mydic={}
        result=[]
        n=0
        for key,value in dic.items():
            n+=1
            if n<=20:
                mydic[key]=value
            else:
                break
        for key,value in mydic.items():
            result.append([key,value])
        c.add("",result) 
        c.set_global_opts(title_opts=opts.TitleOpts(title="专利数量前20位申请人分析饼状图",pos_left="center", pos_top="top"), legend_opts=opts.LegendOpts(type_='scroll',pos_top="bottom"))
        c.set_series_opts(label_opts=opts.LabelOpts(formatter="{b}: {c}"))  
        return c.dump_options_with_quotes()


    def authorbar(self,dic):#绘制柱状图
        c =Bar(init_opts=opts.InitOpts(theme=ThemeType.WALDEN))
        X, Y = [], []
        n=0
        for key,value in dic.items():
            n+=1
            if n<=20:
                X.append(key)
                Y.append(int(value))
            else:
                break
        c.add_xaxis(X)  
        c.add_yaxis(series_name="申请人", y_axis=Y) 
        c.set_global_opts(title_opts=opts.TitleOpts(title="专利数量前20位申请人分析柱状图",pos_left="center", pos_top="top"), legend_opts=opts.LegendOpts(type_='scroll',pos_top="bottom"))       
        return c.dump_options_with_quotes()
        
    def authorlist(self):#将数据中所有作者解析
        dic={}
        for i in self.data:
            mystr=i[1]
            list=str(mystr).split(";")
            for h in list:
                if h[0]==' ':
                    h=h[1:]
                dic[h]=0
        return dic
    
    def listsort(self,dic,id):#给字典赋值并排序
        total=0
        for i in dic.keys():
            i=str(i)
            self.db.myexecu("select count(inventor_list) as cnt from( select inventor_list from patent.patent where  (inventor_list like'%; "+i+"'"+" or inventor_list like'%"+i+";%'"+" or inventor_list like '%;"+i+"'"+" or inventor_list ="+"'"+i+"'"+") and id in"+id+" ) as ccc")
            num=str(self.db.data)
            num=num[2:-4]
            num=int(num)
            dic[i]=num
        for i in dic.values():
            total+=i
        self.total=total
        a1=sorted(dic.items(),key=lambda x:x[1],reverse = True)
        dic=dict(a1)
        self.dic=dic
        return dic

  
class areaanaly(cladata):#地域分析
    

    def certainarea(self,name,id):#某个省份的信息
        self.db.myexecu("SELECT  * FROM patent.patent where id in "+id+" and ( substring(applicant_area,1,INSTR(applicant_area,'(')-1)='"+name+"' or substring(applicant_area,1,INSTR(applicant_area,';')-1)='"+name+"')")
        return self.db.data
        
    def newdic(self):#创建一个字典
        dic={}
        for i in self.data:
            dic[i[1]]=i[0]
        return dic

    def areabar(self,dic):#绘制柱状图
        c =Bar(init_opts=opts.InitOpts(theme=ThemeType.WALDEN))
        X, Y = [], []
        for key,value in dic.items():
                X.append(key)
                Y.append(int(value))
        c.add_xaxis(X)  
        c.add_yaxis(series_name="地区", y_axis=Y) 
        c.set_global_opts(title_opts=opts.TitleOpts(title="地域分析柱状图",pos_left="center", pos_top="top"), legend_opts=opts.LegendOpts(type_='scroll',pos_top="bottom"))   
        return c.dump_options_with_quotes()   

    def areapie(self,dic):#绘制饼状图
        c =Pie(init_opts=opts.InitOpts(theme=ThemeType.WALDEN))
        result=[]
        for key,value in dic.items():
            result.append([key,value])
        c.add("",result) 
        c.set_global_opts(title_opts=opts.TitleOpts(title="地域分析饼状图",pos_left="center", pos_top="top"), legend_opts=opts.LegendOpts(type_='scroll',pos_top="bottom"))
        c.set_series_opts(label_opts=opts.LabelOpts(formatter="{b}: {c}"))           
        return c.dump_options_with_quotes()


class trendanaly(cladata):#趋势分析


            
    def opendic(self,stime,etime):#每遇到一个新类型就新建一个新字典
        dic={}
        i=stime
        while i<=etime:
            dic[i]=0
            i+=1
        return dic
    
    def newdic(self,stime,etime):#创建值为字典的字典对
        updic={}
        for h in self.data:
            updic[h[0]]=self.opendic(stime,etime)               
        for h in self.data:
            updic[h[0]][float(h[1])]=h[2]    
        return updic
               
    def trendpie(self,dic):#绘制每个类别的饼状图
        output=[]
        for key,value in dic.items():
            output.append(self.pietrend(value,key))
        return output
            

    def pietrend(self,dic,title):#绘制饼状图
        c =Pie(init_opts=opts.InitOpts(theme=ThemeType.WALDEN))
        result=[]

        for key,value in dic.items():
            if value != 0:
                result.append([str(key)+"年",value,]) 
        c.add("",result) 
        c.set_global_opts(title_opts=opts.TitleOpts(title="分类号"+str(title)+"趋势分析饼状图",pos_left="center", pos_top="top"), legend_opts=opts.LegendOpts(type_='scroll',pos_top="bottom"))
        c.set_series_opts(label_opts=opts.LabelOpts(formatter="{b}: {c}"))             
        return c.dump_options_with_quotes()
                
    def trendline(self,dic,stime,etime):#绘制折线图
        mytimelist=[]
        mytime=stime
        while mytime<=etime:
            mytimelist.append(str(mytime)+"年")
            mytime+=1
        c =Line(init_opts=opts.InitOpts(theme=ThemeType.WALDEN))      
        c.add_xaxis(mytimelist)
        for key,value in dic.items():
            Y=[]
            for upvalue in value.values():
                Y.append(upvalue)
            c.add_yaxis( series_name=key,y_axis=Y, is_connect_nones=True)
            print(Y)
        c.set_global_opts(title_opts=opts.TitleOpts(title=str(stime)+"年至"+str(etime)+"年趋势分析折线图",pos_left="center", pos_top="top"), legend_opts=opts.LegendOpts(type_='scroll',pos_top="bottom"))
        return c.dump_options_with_quotes()

                
    def trendbar(self,dic,stime,etime):#绘制时间趋势柱状图     
        c =Bar(init_opts=opts.InitOpts(theme=ThemeType.WALDEN))
        mytimelist=[]
        mytime=stime
        while mytime<=etime:
            mytimelist.append(str(mytime)+"年")
            mytime+=1
        c.add_xaxis(mytimelist)    
        mydic={}
        for key,value in dic.items():
            for uvalue in value.values():
                if uvalue !=0:
                    mydic[key]=value
                    break
        for upkey,upvalue in mydic.items():
            Y = []
            for value in upvalue.values():
                Y.append(int(value)) 
            c.add_yaxis( series_name=upkey,y_axis=Y)
        c.set_global_opts(title_opts=opts.TitleOpts(title=str(stime)+"年至"+str(etime)+"年趋势分析柱状图",pos_left="center", pos_top="top"), legend_opts=opts.LegendOpts(type_='scroll',pos_top="bottom"))             
        return c.dump_options_with_quotes()

class pdfcreator:#pdf生成器
    
    def __init__(self,id,db):
        self.id=id
        self.db=db
        self.db.myexecu("SELECT item_type,corr_id FROM patent.report_content_item where report_id="+str(id))
        self.data=self.db.data
        yaml_path= os.path.dirname(os.path.dirname(__file__)) + "/config/config.yml"
        with open(yaml_path,"r",encoding="utf-8") as f:
            data=yaml.load(f,Loader=yaml.FullLoader)
        self.path=data['pdf-output-path']

        
    
    def finalword(self):#生成最终的word  
        a=1
        pagenum={1:"一、",2:"二、",3:"三、"}      
        document = Document()
        mei=document.sections[0].header.paragraphs[0]
        m=document.sections[0].header
        text=mei.add_run("专利分析报告") 
        text.font.size = Pt(10)                
        text.bold = True                    
        text.font.name = 'Arial'           
        text.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')  
        mei.alignment = WD_ALIGN_PARAGRAPH.CENTER 
        m.add_paragraph("-----------------------------------------------------------------------------").paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p=document.add_paragraph("专利分析报告") 
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].font.size = Pt(36)
        p.runs[0].font.name = '宋体'
        p.runs[0].element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        p=document.add_paragraph("目录") 
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].font.size = Pt(16)
        p.runs[0].font.name = '黑体'
        p.runs[0].element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        if(len(self.addsearchresult())!=0):
            pp=document.add_paragraph(pagenum[a]+"专利检索结果") 
            pp.alignment = WD_ALIGN_PARAGRAPH.LEFT
            pp.runs[0].font.size = Pt(12)
            pp.runs[0].font.name = '宋体'
            pp.runs[0].element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            a+=1
        if(len(self.addnovelresult())!=0):
            aa=document.add_paragraph(pagenum[a]+"新颖性分析结果") 
            aa.alignment = WD_ALIGN_PARAGRAPH.LEFT
            aa.runs[0].font.size = Pt(12)
            aa.runs[0].font.name = '宋体'
            aa.runs[0].element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')   
            a+=1
        if(len(self.addpicresult())!=0):
            bb=document.add_paragraph(pagenum[a]+"统计分析结果") 
            bb.alignment = WD_ALIGN_PARAGRAPH.LEFT
            bb.runs[0].font.size = Pt(12)
            bb.runs[0].font.name = '宋体'
            bb.runs[0].element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')   
            a+=1
        document.paragraphs[a].runs[0].add_break(docx.enum.text.WD_BREAK.PAGE)
        ji=1
        npagenum={1:"1、",2:"2、",3:"3、"}
        if(len(self.addsearchresult())!=0):
            aa=document.add_paragraph(npagenum[ji]+"专利检索结果") 
            aa.alignment = WD_ALIGN_PARAGRAPH.CENTER
            aa.runs[0].font.size = Pt(16)
            aa.runs[0].font.name = '黑体'
            aa.runs[0].element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
            for a in self.addsearchresult():
                table = document.add_table(rows=1, cols=3,style='TableGrid')
                cells = table.add_row().cells
                cells[0].text = '标题'
                cells[1].text = '申请人列表'
                cells[2].text = "申请时间"
                rows = table.rows[0]
                for cell in rows.cells:
                    shading_elm = parse_xml(r'<w:shd {} w:fill="D9D9D9"/>'.format(nsdecls('w')))
                    cell._tc.get_or_add_tcPr().append(shading_elm)
                for b in a:
                    for c in b:
                        document.styles['Normal'].font.name = u'宋体'
                        document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
                        cells = table.add_row().cells
                        cells[0].text=c[0]
                        cells[1].text=c[1]
                        cells[2].text=c[2]
            ji+=1  
            document.add_paragraph(" ")      
        if(len(self.addnovelresult())!=0):
            aa=document.add_paragraph(npagenum[ji]+"新颖性分析结果") 
            aa.alignment = WD_ALIGN_PARAGRAPH.CENTER
            aa.runs[0].font.size = Pt(16)
            aa.runs[0].font.name = '黑体'
            aa.runs[0].element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')           
            for a in self.addnovelresult():
                p=document.add_paragraph()
                for b in a:
                    (key, value), = b.items()
                    if key=="有以下相关主权项：":
                        document.styles['Normal'].font.name = u'宋体'
                        document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
                        p.add_run(key)
                        p.add_run("\n")
                    else:
                        document.styles['Normal'].font.name = u'宋体'
                        document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
                        p.add_run(key)                    
                        p.add_run(value)
                        p.add_run("______________________________________________________________________________")
                        p.add_run("\n")
                        #document.add_paragraph("______________________________________________________________________________")
            ji+=1
            document.add_paragraph(" ")
        if(len(self.addpicresult())!=0):
            aa=document.add_paragraph(npagenum[ji]+"统计分析结果") 
            aa.alignment = WD_ALIGN_PARAGRAPH.CENTER
            aa.runs[0].font.size = Pt(16)
            aa.runs[0].font.name = '黑体'
            aa.runs[0].element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')   
            for i in self.addpicresult():
                if i["type"]=="柱状":
                    c =Bar(init_opts=opts.InitOpts(theme=ThemeType.WALDEN))
                    c.add_xaxis(i["xaxis"]) 
                    if len(i["data"])==1: 
                        (key, value), = i["data"].items()
                        c.add_yaxis( series_name=key,y_axis=value)
                    else:          
                        for key,value in i["data"].items():  
                            c.add_yaxis( series_name=key,y_axis=value)
                    c.set_global_opts(title_opts=opts.TitleOpts(title=i["title"],pos_left="center", pos_top="top"), legend_opts=opts.LegendOpts(type_='plain',pos_top="bottom")) 
                    make_snapshot(snapshot, c.render(), self.path+"\\"+str(self.id)+"bar.png",is_remove_html=True)
                    paragraph=document.add_paragraph()
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT     
                    run = paragraph.add_run("")
                    run.add_picture(self.path+"\\"+str(self.id)+"bar.png",width=Inches(7.0))
                    document.add_paragraph("   ")
                if i["type"]=="折线":
                    c =Line(init_opts=opts.InitOpts(theme=ThemeType.WALDEN))      
                    c.add_xaxis(i["xaxis"])
                    if len(i["data"])==1: 
                        (key, value), = i["data"].items()
                        c.add_yaxis( series_name=key,y_axis=value,is_connect_nones=True)
                    else:          
                        for key,value in i["data"].items():  
                            c.add_yaxis( series_name=key,y_axis=value,is_connect_nones=True)
                    c.set_global_opts(title_opts=opts.TitleOpts(title=i["title"],pos_left="center", pos_top="top"), legend_opts=opts.LegendOpts(type_='plain',pos_top="bottom"))
                    make_snapshot(snapshot, c.render(), self.path+"\\"+str(self.id)+"line.png",is_remove_html=True)
                    paragraph=document.add_paragraph()
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT     
                    run = paragraph.add_run("")
                    run.add_picture(self.path+"\\"+str(self.id)+"line.png",width=Inches(7.0))
                    document.add_paragraph("   ")
                if i["type"]=="饼状":
                    c =Pie(init_opts=opts.InitOpts(theme=ThemeType.WALDEN))
                    c.add("",i["data"]) 
                    c.set_global_opts(title_opts=opts.TitleOpts(title=i["title"],pos_left="center", pos_top="top"), legend_opts=opts.LegendOpts(type_='plain',pos_top="bottom"))
                    c.set_series_opts(label_opts=opts.LabelOpts(formatter="{b}: {c}"))  
                    make_snapshot(snapshot, c.render(), self.path+"\\"+str(self.id)+"pie.png",is_remove_html=True)
                    paragraph=document.add_paragraph()
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT    
                    run = paragraph.add_run("")
                    run.add_picture(self.path+"\\"+str(self.id)+"pie.png",width=Inches(7.0))
                    document.add_paragraph("   ")
        document.save(self.path+"\\"+str(self.id)+".docx")
    
    def addsearchresult(self):#专利检索内容增加
        tabledata=[]
        mydata=[]
        for i in self.data:            
            if i[0]=="检索结果":
                updata=[]
                self.db.myexecu("SELECT patent_id FROM patent.search_result_item where search_result_id="+str(i[1]))
                for a in self.db.data:
                    mydata.append(a[0])
                for a in mydata:
                    uppdata=[]
                    self.db.myexecu("SELECT title,inventor_list,application_date FROM patent.patent where id ="+str(a))
                    for b in self.db.data:
                        uppdata.append(b)
                    updata.append(uppdata)
                tabledata.append(updata)
        return tabledata
    
    def addnovelresult(self):#新颖性内容增加
        mydata=[]
        for i in self.data:
            if i[0]=="新颖性比对结果":
                updata=[]
                self.db.myexecu("SELECT ori_signory FROM patent.novelty_ana_result where novelty_ana_id="+str(i[1]))
                updata.append({"原主权项：":self.db.data[0]})
                updata.append({"有以下相关主权项：":""})
                self.db.myexecu("SELECT relevant_sig,compare_result,ori_patent_title FROM patent.novelty_ana_item where novelty_ana_id="+str(i[1]))
                for a in self.db.data:
                    updata.append({"相关主权项：":a[0]})
                    updata.append({"来自专利：":a[2]}) 
                    updata.append({"审查意见为：":a[1]})                   
                mydata.append(updata)
        return mydata
    
    def addpicresult(self):#图片内容
        outresult=[]
        tabledata=[]
        for i in self.data:
            if i[0]=="统计分析结果":
                self.db.myexecu("SELECT option_json FROM patent.stats_ana_item where stats_ana_id="+str(i[1]))
                myjson=json.loads(str(self.db.data[0])[2:-3])
                mytitle=myjson["title"][0]["text"]
                type=mytitle[-3:-1]
                if type=="柱状":
                    barresult={}
                    barresult["type"]="柱状"
                    barresult["title"]=mytitle
                    time=[]
                    for a in myjson["xAxis"][0]["data"]:
                        time.append(a)
                    barresult["xaxis"]=time
                    mydata={}
                    for a in myjson["series"]:
                        mydata[a["name"]]=a["data"]
                    barresult["data"]=mydata
                    outresult.append(barresult)
                if type=="折线":
                    barresult={}
                    barresult["type"]="折线"
                    barresult["title"]=mytitle
                    time=[]
                    for a in myjson["xAxis"][0]["data"]:
                        time.append(a)
                    barresult["xaxis"]=time
                    mydata={}
                    for a in myjson["series"]:
                        updata=[]
                        for b in a["data"]:
                            updata.append(b[1])
                        mydata[a["name"]]=updata
                    barresult["data"]=mydata
                    outresult.append(barresult)
                if type=="饼状":
                    barresult={}
                    barresult["type"]="饼状"
                    barresult["title"]=mytitle
                    time=[]
                    mydata=[]
                    a = myjson["series"][0]["data"]
                    for b in a:
                        mydata.append([b["name"],b["value"]])

                    barresult["data"]=mydata
                    outresult.append(barresult)
        return outresult
        

  
    
    def pdfcreate(self):#生成最后的PDF
        convert(self.path+"\\"+str(self.id)+".docx", self.path+"\\"+str(self.id)+".pdf")
        addstr=self.path+"\\"+"\\"+str(self.id)+".pdf"
        return addstr



    
def myinput(list): #将系统的输入转换成sql语句中所要的格式
    mystr='('
    for i in list:
        mystr+=str(i)
        mystr+=','
    mystr=mystr[0:-1]
    mystr+=')'
    return mystr


def trendsql(stime,etime,list):#趋势分析sql语句
    sqlstr="SELECT substring(main_class_list,1,INSTR(main_class_list,'/')-1) as clss_type, substring(publication_date,1,4) as pub_year,count(patent_type) as cnt FROM patent.patent where substring(publication_date,1,4)<="+str(etime)+" "+"and substring(publication_date,1,4)>="+str(stime)+" "+"and id in"+myinput(list)+" "+"group by substring(main_class_list,1,INSTR(main_class_list,'/')-1),pub_year"      
    return sqlstr


def authorsql(list):#发明人分析sql语句
    sqlstr="select id,inventor_list from patent.patent where id in "+myinput(list)
    return sqlstr


def areasql(list):#地域分析sql语句
    sqlstr="select count(c2),c2 from(SELECT  applicant_area,CONCAT(substring(applicant_area,1,INSTR(applicant_area,';')-1) ,substring(applicant_area,1,INSTR(applicant_area,'(')-1))  as c2 FROM patent.patent where id in "+myinput(list)+" ) as ttt group by c2"
    return sqlstr


def timecal(data):#时间置信区间计算
    num=[]
    for i in data:
        num.append(int(i[1]))
    result=st.norm.interval(alpha=0.99,loc=np.mean(num),scale=st.sem(num))
    output=[]
    if int(result[0]) == int(result[1]):
        output.append(int(result[0])-1)
        output.append(int(result[1])+1)
    else:
        output.append(int(result[0]))
        output.append(int(result[1]))
    return output
    
    
def timesql(list):#置信区间计算SQL语句
    sqlstr="SELECT  id,substring(publication_date,1,4) as pub_year FROM patent.patent where id in "+myinput(list)
    return sqlstr

# #申请人分析
def author(list,type):
    db = mydb('152.136.114.189','zym','zym','patent',6336,'utf8')
    output=[]    
    sqls=authorsql(list)
    db.myexecu(sqls)
    dd=authoranaly(db,'234')
    dic=dd.authorlist()
    dic=dd.listsort(dic,myinput(list))
    if type == "bar":
       output.append(dd.authorbar(dic))
       db.endconn()
       return output     
    output.append(dd.authorpie(dic,'234'))
    db.endconn() 
    return output

# #趋势分析
def trend(list,type):
    db = mydb('152.136.114.189','zym','zym','patent',6336,'utf8')
    sqls=timesql(list)
    db.myexecu(sqls)
    timeblock=timecal(db.data)
    output=[]
    sqls=trendsql(timeblock[0],timeblock[1],list)
    db.myexecu(sqls)
    dd=trendanaly(db,"234")
    dic=dd.newdic(timeblock[0],timeblock[1])
    if type == "bar":
        output.append(dd.trendbar(dic,timeblock[0],timeblock[1]))
        return output
    if type == "line":
        output.append(dd.trendline(dic,timeblock[0],timeblock[1]))
        return output    
    for i in dd.trendpie(dic):
        output.append(i)
    db.endconn()
    return output
  
# #地域分析
def area(list,type):
    db = mydb('152.136.114.189','zym','zym','patent',6336,'utf8')
    output=[]
    sqls=areasql(list)
    db.myexecu(sqls)
    dd=areaanaly(db,'234')
    dic=dd.newdic()
    if type == "pie":
        output.append(dd.areapie(dic))
        db.endconn()
        return output
    output.append(dd.areabar(dic))
    db.endconn()
    return output

def analyze_by_list(patentIds, figType, anaType):#三个功能封装在一起
    if anaType == 'author':
        return author(patentIds,figType)
    if anaType == 'area':
        return area(patentIds,figType)
    return trend(patentIds,figType)

def pdf_output(id):#输出pdf
    db = mydb('152.136.114.189','zym','zym','patent',6336,'utf8')
    mypdf=pdfcreator(id,db)
    mypdf.finalword()
    output=mypdf.pdfcreate()
    db.myexecu('update patent.report2generate set status="已生成",pdf_file_path="'+output+'" where report_id='+str(mypdf.id))
    db.endconn()
    print("报告生成完毕")
    return  









